"""Tests for PDF and Word document reading in ReadFileTool.

Tests are organized into:
- TestExtractorModule: Direct unit tests for document_extractor.py functions
- TestReadPDF: Integration tests for PDF reading via ReadFileTool (subprocess)
- TestReadDOCX: Integration tests for DOCX reading via ReadFileTool (subprocess)
- TestDocumentSecurity: Security guards (file size, zip bomb, line cap, crash isolation)
- TestDocumentEdgeCases: Edge cases and metadata
- TestMultimodalExtraction: Image extraction from PDF/DOCX
- TestMultimodalToolResult: Multimodal content building in ReadFileTool
- TestFrameToolResultMultimodal: _frame_tool_result with str and list content
- TestAnthropicMultimodalToolResult: Anthropic backend multimodal tool results
"""

import importlib.util
import json
import subprocess
import zipfile

import pytest
from pathlib import Path
from unittest.mock import patch, MagicMock

# Prime import chain (see conftest.py)
import src.core  # noqa: F401

from src.tools import ReadFileTool
from src.tools.file_operations import FileOperationTool
from src.tools.base import ToolStatus
from src.tools.document_extractor import extract_pdf, extract_docx, check_zip_bomb, _parse_pages
from src.llm.base import LLMConfig, LLMBackendType


@pytest.fixture(autouse=True)
def allow_test_workspace(tmp_path, monkeypatch):
    """Allow file operations in test tmp_path."""
    monkeypatch.setattr(FileOperationTool, "_workspace_root", tmp_path)
    yield
    monkeypatch.setattr(FileOperationTool, "_workspace_root", None)


def _make_subprocess_result(lines, error=None, returncode=0, images=None):
    """Create a mock subprocess.CompletedProcess with JSON output."""
    data = {"lines": lines, "images": images or [], "error": error}
    result = MagicMock(spec=subprocess.CompletedProcess)
    result.stdout = json.dumps(data)
    result.stderr = ""
    result.returncode = returncode
    return result


# ---------------------------------------------------------------------------
# Extractor module unit tests (no subprocess, direct function calls)
# ---------------------------------------------------------------------------


_has_fitz = bool(importlib.util.find_spec("fitz"))


@pytest.mark.skipif(not _has_fitz, reason="PyMuPDF (fitz) not installed")
class TestExtractorModule:
    """Direct tests for document_extractor.py extraction functions."""

    def test_extract_pdf_basic(self, tmp_path):
        """Test PDF extraction produces page headers and text."""
        import fitz

        # Create a real minimal PDF with text
        pdf_path = tmp_path / "test.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Hello PDF World")
        doc.save(str(pdf_path))
        doc.close()

        lines, images, error = extract_pdf(pdf_path, max_lines=10_000)

        assert error is None
        assert images == []
        assert any("--- Page 1 of 1 ---" in line for line in lines)
        assert any("Hello PDF World" in line for line in lines)

    def test_extract_pdf_line_cap(self, tmp_path):
        """Test PDF extraction respects max_lines limit."""
        import fitz

        pdf_path = tmp_path / "big.pdf"
        doc = fitz.open()
        page = doc.new_page()
        # Insert many lines of text
        for i in range(100):
            page.insert_text((72, 72 + i * 8), f"Line {i}")
        doc.save(str(pdf_path))
        doc.close()

        lines, _images, error = extract_pdf(pdf_path, max_lines=10)

        assert error is None
        assert len(lines) == 11  # 10 lines + truncation notice
        assert "extraction stopped" in lines[-1]

    def test_extract_docx_basic(self, tmp_path):
        """Test DOCX extraction produces paragraph text."""
        from docx import Document

        docx_path = tmp_path / "test.docx"
        doc = Document()
        doc.add_paragraph("Hello")
        doc.add_paragraph("World")
        doc.save(str(docx_path))

        lines, images, error = extract_docx(docx_path, max_lines=10_000,
                                             max_zip_size=200 * 1024 * 1024, max_zip_ratio=100)

        assert error is None
        assert images == []
        assert "Hello" in lines
        assert "World" in lines

    def test_extract_docx_with_table(self, tmp_path):
        """Test DOCX extraction includes tables."""
        from docx import Document

        docx_path = tmp_path / "table.docx"
        doc = Document()
        doc.add_paragraph("Before")
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[1].cells[0].text = "C"
        table.rows[1].cells[1].text = "D"
        doc.save(str(docx_path))

        lines, _images, error = extract_docx(docx_path, max_lines=10_000,
                                              max_zip_size=200 * 1024 * 1024, max_zip_ratio=100)

        assert error is None
        assert "Before" in lines
        assert any("[Table 1]" in line for line in lines)
        assert any("| A | B |" in line for line in lines)

    def test_extract_docx_line_cap(self, tmp_path):
        """Test DOCX extraction respects max_lines limit."""
        from docx import Document

        docx_path = tmp_path / "big.docx"
        doc = Document()
        for i in range(50):
            doc.add_paragraph(f"Paragraph {i}")
        doc.save(str(docx_path))

        lines, _images, error = extract_docx(docx_path, max_lines=10,
                                              max_zip_size=200 * 1024 * 1024, max_zip_ratio=100)

        assert error is None
        assert len(lines) == 11  # 10 lines + truncation notice
        assert "extraction stopped" in lines[-1]

    def test_check_zip_bomb_normal_file(self, tmp_path):
        """Test that normal DOCX passes zip bomb check."""
        from docx import Document

        docx_path = tmp_path / "normal.docx"
        doc = Document()
        doc.add_paragraph("Normal content")
        doc.save(str(docx_path))

        result = check_zip_bomb(docx_path)
        assert result is None

    def test_check_zip_bomb_high_ratio(self, tmp_path):
        """Test that suspicious compression ratio is detected."""
        # Create a ZIP with artificially high compression ratio
        zip_path = tmp_path / "bomb.docx"
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            # Write highly compressible data (1MB of zeros)
            zf.writestr("word/document.xml", "\x00" * (1024 * 1024))

        result = check_zip_bomb(zip_path, max_ratio=5)
        assert result is not None
        assert "compression ratio" in result.lower() or "Suspicious" in result


# ---------------------------------------------------------------------------
# PDF integration tests (ReadFileTool -> subprocess)
# ---------------------------------------------------------------------------


class TestReadPDF:
    """Integration tests for PDF reading via ReadFileTool subprocess."""

    def test_pdf_basic_text(self, tmp_path):
        """Test PDF text extraction via subprocess."""
        pdf_path = tmp_path / "test.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result(
            ["--- Page 1 of 1 ---", "", "Line one", "Line two", "Line three"]
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.SUCCESS
        assert "Line one" in result.output
        assert "Line two" in result.output
        assert "--- Page 1 of 1 ---" in result.output

    def test_pdf_multi_page(self, tmp_path):
        """Test multi-page PDF output."""
        pdf_path = tmp_path / "multi.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result([
            "--- Page 1 of 3 ---", "", "Page one content", "",
            "--- Page 2 of 3 ---", "", "Page two content", "",
            "--- Page 3 of 3 ---", "", "Page three content", "",
        ])

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.SUCCESS
        assert "--- Page 1 of 3 ---" in result.output
        assert "--- Page 3 of 3 ---" in result.output

    def test_pdf_with_tables(self, tmp_path):
        """Test PDF table extraction output."""
        pdf_path = tmp_path / "tables.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result([
            "--- Page 1 of 1 ---", "", "Some text", "",
            "[Table 1]", "| Name | Age |", "| Alice | 30 |",
        ])

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.SUCCESS
        assert "[Table 1]" in result.output
        assert "| Name | Age |" in result.output
        assert "| Alice | 30 |" in result.output

    def test_pdf_line_range(self, tmp_path):
        """Test line range selection on PDF output."""
        pdf_path = tmp_path / "range.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result(
            [f"Line {i}" for i in range(20)]
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path), start_line=3, max_lines=2)

        assert result.status == ToolStatus.SUCCESS
        assert result.metadata["lines_returned"] == 2
        assert result.metadata["start_line"] == 3
        assert result.metadata["has_more"] is True

    def test_pdf_extractor_returns_error(self, tmp_path):
        """Test handling when extractor subprocess reports an error."""
        pdf_path = tmp_path / "bad.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result(
            [], error="PyMuPDF is required to read PDF files. Install it: pip install PyMuPDF"
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.ERROR
        assert "PyMuPDF" in result.error


# ---------------------------------------------------------------------------
# DOCX integration tests (real files through subprocess)
# ---------------------------------------------------------------------------


class TestReadDOCX:
    """Integration tests for Word document reading."""

    def _create_docx_with_text(self, path: Path, paragraphs: list[str]):
        """Create a real .docx file with given paragraphs."""
        from docx import Document

        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)
        doc.save(str(path))

    def _create_docx_with_table(self, path: Path, headers: list[str], rows: list[list[str]]):
        """Create a real .docx file with a paragraph and a table."""
        from docx import Document

        doc = Document()
        doc.add_paragraph("Text before table")
        table = doc.add_table(rows=1, cols=len(headers))
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
        for row_data in rows:
            row = table.add_row()
            for i, cell in enumerate(row_data):
                row.cells[i].text = cell
        doc.add_paragraph("Text after table")
        doc.save(str(path))

    def test_docx_basic_text(self, tmp_path):
        """Test extracting paragraphs from a Word document."""
        docx_path = tmp_path / "test.docx"
        self._create_docx_with_text(docx_path, ["Hello", "World", "Test"])

        tool = ReadFileTool()
        result = tool.execute(file_path=str(docx_path))

        assert result.status == ToolStatus.SUCCESS
        assert "Hello" in result.output
        assert "World" in result.output
        assert "Test" in result.output

    def test_docx_with_table(self, tmp_path):
        """Test extracting a table from a Word document."""
        docx_path = tmp_path / "table.docx"
        self._create_docx_with_table(
            docx_path,
            headers=["Name", "Score"],
            rows=[["Alice", "95"], ["Bob", "87"]],
        )

        tool = ReadFileTool()
        result = tool.execute(file_path=str(docx_path))

        assert result.status == ToolStatus.SUCCESS
        assert "[Table 1]" in result.output
        assert "| Name | Score |" in result.output
        assert "Text before table" in result.output

    def test_docx_preserves_order(self, tmp_path):
        """Test that paragraphs and tables appear in document order."""
        from docx import Document

        docx_path = tmp_path / "order.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        table = doc.add_table(rows=1, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        doc.add_paragraph("Second paragraph")
        doc.save(str(docx_path))

        tool = ReadFileTool()
        result = tool.execute(file_path=str(docx_path))

        assert result.status == ToolStatus.SUCCESS
        lines = result.output
        first_pos = lines.find("First paragraph")
        table_pos = lines.find("[Table 1]")
        second_pos = lines.find("Second paragraph")
        assert first_pos < table_pos < second_pos

    def test_docx_empty_document(self, tmp_path):
        """Test reading an empty Word document."""
        from docx import Document

        docx_path = tmp_path / "empty.docx"
        doc = Document()
        doc.save(str(docx_path))

        tool = ReadFileTool()
        result = tool.execute(file_path=str(docx_path))

        assert result.status == ToolStatus.SUCCESS

    def test_docx_line_range(self, tmp_path):
        """Test that start_line/max_lines work on DOCX output."""
        docx_path = tmp_path / "range.docx"
        self._create_docx_with_text(docx_path, [f"Para {i}" for i in range(20)])

        tool = ReadFileTool()
        result = tool.execute(file_path=str(docx_path), start_line=5, max_lines=3)

        assert result.status == ToolStatus.SUCCESS
        assert result.metadata["lines_returned"] == 3
        assert result.metadata["start_line"] == 5
        assert result.metadata["has_more"] is True


# ---------------------------------------------------------------------------
# Security tests
# ---------------------------------------------------------------------------


class TestDocumentSecurity:
    """Security guards: file size, zip bomb, line cap, crash isolation."""

    def test_file_size_limit_rejects_large_pdf(self, tmp_path):
        """Test that PDFs exceeding MAX_DOCUMENT_SIZE_BYTES are rejected."""
        pdf_path = tmp_path / "huge.pdf"
        pdf_path.write_bytes(b"x" * 1000)

        tool = ReadFileTool()
        tool.MAX_DOCUMENT_SIZE_BYTES = 500  # Set low for testing

        result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.ERROR
        assert "too large" in result.error

    def test_file_size_limit_rejects_large_docx(self, tmp_path):
        """Test that DOCX exceeding MAX_DOCUMENT_SIZE_BYTES are rejected."""
        docx_path = tmp_path / "huge.docx"
        docx_path.write_bytes(b"x" * 1000)

        tool = ReadFileTool()
        tool.MAX_DOCUMENT_SIZE_BYTES = 500

        result = tool.execute(file_path=str(docx_path))

        assert result.status == ToolStatus.ERROR
        assert "too large" in result.error

    def test_file_size_limit_allows_small_files(self, tmp_path):
        """Test that files under the limit are processed normally."""
        from docx import Document

        docx_path = tmp_path / "small.docx"
        doc = Document()
        doc.add_paragraph("Small file")
        doc.save(str(docx_path))

        tool = ReadFileTool()
        # Default 10MB limit should easily allow this
        result = tool.execute(file_path=str(docx_path))

        assert result.status == ToolStatus.SUCCESS

    def test_zip_bomb_detected(self, tmp_path):
        """Test that DOCX zip bombs are rejected by the extractor."""
        zip_path = tmp_path / "bomb.docx"
        with zipfile.ZipFile(str(zip_path), "w", zipfile.ZIP_DEFLATED) as zf:
            # Write highly compressible data
            zf.writestr("word/document.xml", "\x00" * (2 * 1024 * 1024))

        # Use the extractor directly to test the zip bomb guard
        lines, _images, error = extract_docx(
            zip_path, max_lines=10_000,
            max_zip_size=200 * 1024 * 1024, max_zip_ratio=5  # Very strict ratio
        )

        assert error is not None
        assert "compression ratio" in error.lower() or "Suspicious" in error

    def test_subprocess_crash_returns_error(self, tmp_path):
        """Test that a subprocess crash (e.g. segfault) returns an error, not a crash."""
        pdf_path = tmp_path / "crash.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = MagicMock(spec=subprocess.CompletedProcess)
        mock_result.stdout = ""
        mock_result.stderr = ""
        mock_result.returncode = -11  # SIGSEGV

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.ERROR
        assert "corrupt or unsupported" in result.error

    def test_subprocess_timeout_returns_error(self, tmp_path):
        """Test that a subprocess timeout returns an error gracefully."""
        pdf_path = tmp_path / "slow.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        tool = ReadFileTool()
        with patch(
            "src.tools.file_operations.subprocess.run",
            side_effect=subprocess.TimeoutExpired(cmd="test", timeout=30),
        ):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.ERROR
        assert "timed out" in result.error

    def test_subprocess_bad_json_returns_error(self, tmp_path):
        """Test that garbled subprocess output returns an error."""
        pdf_path = tmp_path / "garbled.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = MagicMock(spec=subprocess.CompletedProcess)
        mock_result.stdout = "NOT JSON {{{{"
        mock_result.stderr = ""
        mock_result.returncode = 0

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.ERROR

    def test_error_message_sanitized(self, tmp_path):
        """Test that internal error details are not leaked to the LLM."""
        pdf_path = tmp_path / "err.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = MagicMock(spec=subprocess.CompletedProcess)
        mock_result.stdout = ""
        mock_result.stderr = ""
        mock_result.returncode = 1  # Non-zero = crash

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.ERROR
        # Should show generic message, not internal paths or tracebacks
        assert "corrupt or unsupported" in result.error

    def test_corrupt_pdf_via_subprocess(self, tmp_path):
        """Test that a truly corrupt PDF file is handled gracefully end-to-end."""
        pdf_path = tmp_path / "corrupt.pdf"
        pdf_path.write_bytes(b"this is not a pdf at all")

        # Let it run the real subprocess - the extractor should catch the error
        tool = ReadFileTool()
        result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.ERROR

    def test_corrupt_docx_via_subprocess(self, tmp_path):
        """Test that a truly corrupt DOCX file is handled gracefully end-to-end."""
        docx_path = tmp_path / "corrupt.docx"
        docx_path.write_bytes(b"this is not a docx at all")

        tool = ReadFileTool()
        result = tool.execute(file_path=str(docx_path))

        assert result.status == ToolStatus.ERROR


# ---------------------------------------------------------------------------
# Edge cases
# ---------------------------------------------------------------------------


class TestDocumentEdgeCases:
    """Edge cases shared across document formats."""

    def test_unknown_binary_format_not_intercepted(self, tmp_path):
        """Test that non-PDF/DOCX binary files are NOT routed to extractors."""
        bin_path = tmp_path / "data.bin"
        bin_path.write_bytes(b"\x00\x01\x02\x03")

        tool = ReadFileTool()
        result = tool.execute(file_path=str(bin_path))

        assert result.status == ToolStatus.SUCCESS

    def test_pdf_extension_case_insensitive(self, tmp_path):
        """Test that .PDF (uppercase) is still routed to the extractor."""
        pdf_path = tmp_path / "test.PDF"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result(
            ["--- Page 1 of 1 ---", "", "Upper case extension"]
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.SUCCESS
        assert "Upper case extension" in result.output

    def test_metadata_structure(self, tmp_path):
        """Test that metadata from document reads matches plain text read structure."""
        from docx import Document

        docx_path = tmp_path / "meta.docx"
        doc = Document()
        doc.add_paragraph("Test line")
        doc.save(str(docx_path))

        tool = ReadFileTool()
        result = tool.execute(file_path=str(docx_path))

        assert result.metadata is not None
        required_keys = {"file_path", "total_lines", "lines_returned", "start_line", "end_line", "has_more"}
        assert required_keys.issubset(set(result.metadata.keys()))

    def test_text_file_unaffected(self, tmp_path):
        """Test that plain text files still use the streaming path, not subprocess."""
        txt_path = tmp_path / "plain.txt"
        txt_path.write_text("line 1\nline 2\nline 3\n")

        tool = ReadFileTool()
        result = tool.execute(file_path=str(txt_path))

        assert result.status == ToolStatus.SUCCESS
        assert "line 1" in result.output
        assert result.metadata["total_lines"] == 3


# ---------------------------------------------------------------------------
# Multimodal (image extraction) tests
# ---------------------------------------------------------------------------


@pytest.mark.skipif(not _has_fitz, reason="PyMuPDF (fitz) not installed")
class TestMultimodalExtraction:
    """Tests for scout-then-render image extraction from documents."""

    def test_pdf_scout_mode_shows_hints(self, tmp_path):
        """Test that scout mode (default) shows image hints but no images."""
        import fitz
        import base64

        pdf_path = tmp_path / "with_image.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Page with image")
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
            "nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
        )
        page.insert_image(fitz.Rect(100, 100, 200, 200), stream=tiny_png)
        doc.save(str(pdf_path))
        doc.close()

        lines, images, error = extract_pdf(pdf_path, max_lines=10_000, extract_images=False)

        assert error is None
        assert images == []  # No images in scout mode
        assert any("use extract_images=true to see" in line for line in lines)

    def test_pdf_render_mode_returns_page_screenshots(self, tmp_path):
        """Test that render mode returns page screenshots for pages with images."""
        import fitz
        import base64

        pdf_path = tmp_path / "with_image.pdf"
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((72, 72), "Page with image")
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
            "nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
        )
        page.insert_image(fitz.Rect(100, 100, 200, 200), stream=tiny_png)
        doc.save(str(pdf_path))
        doc.close()

        lines, images, error = extract_pdf(pdf_path, max_lines=10_000, extract_images=True)

        assert error is None
        assert len(images) >= 1
        assert images[0]["media_type"] == "image/png"
        assert len(images[0]["base64"]) > 0
        assert images[0]["source"] == "page_1"
        # Should have [IMAGE:0] marker interleaved in lines
        assert any("[IMAGE:0]" in line for line in lines)

    def test_pdf_text_only_pages_not_rendered(self, tmp_path):
        """Test that text-only pages get no image rendering."""
        import fitz
        import base64

        pdf_path = tmp_path / "mixed.pdf"
        doc = fitz.open()
        # Page 1: text only
        page1 = doc.new_page()
        page1.insert_text((72, 72), "Text only page")
        # Page 2: has image
        page2 = doc.new_page()
        page2.insert_text((72, 72), "Page with image")
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
            "nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
        )
        page2.insert_image(fitz.Rect(100, 100, 200, 200), stream=tiny_png)
        doc.save(str(pdf_path))
        doc.close()

        lines, images, error = extract_pdf(pdf_path, max_lines=10_000, extract_images=True)

        assert error is None
        # Only page 2 should produce a screenshot
        assert len(images) == 1
        assert images[0]["source"] == "page_2"

    def test_pdf_image_count_limit(self, tmp_path):
        """Test that max_images cap is respected."""
        import fitz

        pdf_path = tmp_path / "many_images.pdf"
        doc = fitz.open()
        for i in range(5):
            page = doc.new_page()
            pix = fitz.Pixmap(fitz.csRGB, fitz.IRect(0, 0, 2, 2), 0)
            pix.clear_with(i * 40)
            page.insert_image(fitz.Rect(10, 10, 50, 50), pixmap=pix)
        doc.save(str(pdf_path))
        doc.close()

        lines, images, error = extract_pdf(
            pdf_path, max_lines=10_000, extract_images=True, max_images=2,
        )

        assert error is None
        assert len(images) == 2

    def test_pdf_pages_filter(self, tmp_path):
        """Test that pages parameter filters to specific pages."""
        import fitz

        pdf_path = tmp_path / "multi.pdf"
        doc = fitz.open()
        for i in range(5):
            page = doc.new_page()
            page.insert_text((72, 72), f"Page {i+1} content")
        doc.save(str(pdf_path))
        doc.close()

        # Only read pages 2 and 4
        pages_filter = _parse_pages("2,4", 5)
        lines, images, error = extract_pdf(
            pdf_path, max_lines=10_000, pages_filter=pages_filter,
        )

        assert error is None
        assert any("Page 2 of 5" in line for line in lines)
        assert any("Page 4 of 5" in line for line in lines)
        assert not any("Page 1 of 5" in line for line in lines)
        assert not any("Page 3 of 5" in line for line in lines)

    def test_parse_pages_various_formats(self):
        """Test page specification parsing."""
        assert _parse_pages("3", 10) == {2}
        assert _parse_pages("1-3", 10) == {0, 1, 2}
        assert _parse_pages("2,5,8", 10) == {1, 4, 7}
        assert _parse_pages("1-3,7", 10) == {0, 1, 2, 6}

    def test_docx_scout_mode_shows_hints(self, tmp_path):
        """Test that DOCX scout mode shows image hints."""
        from docx import Document
        from docx.shared import Inches
        import base64

        docx_path = tmp_path / "with_image.docx"
        doc = Document()
        doc.add_paragraph("Document with image")
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
            "nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
        )
        img_path = tmp_path / "tiny.png"
        img_path.write_bytes(tiny_png)
        doc.add_picture(str(img_path), width=Inches(1))
        doc.save(str(docx_path))

        lines, images, error = extract_docx(
            docx_path, max_lines=10_000,
            max_zip_size=200 * 1024 * 1024, max_zip_ratio=100,
            extract_images=False,
        )

        assert error is None
        assert images == []
        assert any("use extract_images=true to see" in line for line in lines)

    def test_docx_render_mode_extracts_inline(self, tmp_path):
        """Test that DOCX render mode extracts images at paragraph position."""
        from docx import Document
        from docx.shared import Inches
        import base64

        docx_path = tmp_path / "with_image.docx"
        doc = Document()
        doc.add_paragraph("Before image")
        tiny_png = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAADUlEQVR4"
            "nGP4z8BQDwAEgAF/pooBPQAAAABJRU5ErkJggg=="
        )
        img_path = tmp_path / "tiny.png"
        img_path.write_bytes(tiny_png)
        doc.add_picture(str(img_path), width=Inches(1))
        doc.add_paragraph("After image")
        doc.save(str(docx_path))

        lines, images, error = extract_docx(
            docx_path, max_lines=10_000,
            max_zip_size=200 * 1024 * 1024, max_zip_ratio=100,
            extract_images=True,
        )

        assert error is None
        assert len(images) >= 1
        assert images[0]["media_type"] == "image/png"
        # Image marker should be between "Before image" and "After image"
        marker_idx = next(i for i, l in enumerate(lines) if "[IMAGE:0]" in l)
        before_idx = next(i for i, l in enumerate(lines) if "Before image" in l)
        after_idx = next(i for i, l in enumerate(lines) if "After image" in l)
        assert before_idx < marker_idx < after_idx


class TestMultimodalToolResult:
    """Tests for multimodal content building with interleaved markers."""

    def test_interleaved_markers_produce_multimodal(self, tmp_path):
        """Test that [IMAGE:N] markers in text are replaced with image blocks."""
        pdf_path = tmp_path / "img.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_images = [
            {"base64": "aGVsbG8=", "media_type": "image/png", "source": "page_1"},
        ]
        mock_result = _make_subprocess_result(
            ["--- Page 1 ---", "Text before image", "[IMAGE:0]", "Text after image"],
            images=mock_images,
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path), extract_images=True)

        assert result.status == ToolStatus.SUCCESS
        assert isinstance(result.output, list)
        # Should have text, image, text, count note
        types = [b["type"] for b in result.output]
        assert "image_url" in types
        assert result.metadata["image_count"] == 1
        # Image should be between text segments, not at end
        img_idx = types.index("image_url")
        assert img_idx > 0  # not first
        assert img_idx < len(types) - 1  # not last (count note is last)

    def test_no_images_returns_string(self, tmp_path):
        """Test that documents without images return plain string output."""
        pdf_path = tmp_path / "text_only.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result(
            ["--- Page 1 ---", "Text only"],
            images=[],
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert result.status == ToolStatus.SUCCESS
        assert isinstance(result.output, str)

    def test_scout_mode_hints_no_images(self, tmp_path):
        """Test that scout mode (default) returns string with hints."""
        pdf_path = tmp_path / "with_hint.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_result = _make_subprocess_result(
            ["--- Page 1 ---", "Text", "[1 image(s) on this page -- use extract_images=true to see]"],
            images=[],
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path))

        assert isinstance(result.output, str)
        assert "extract_images=true" in result.output

    def test_multiple_interleaved_images(self, tmp_path):
        """Test handling of multiple interleaved images."""
        pdf_path = tmp_path / "multi_img.pdf"
        pdf_path.write_bytes(b"%PDF-1.4 dummy")

        mock_images = [
            {"base64": "img1data", "media_type": "image/png", "source": "page_1"},
            {"base64": "img2data", "media_type": "image/jpeg", "source": "page_3"},
        ]
        mock_result = _make_subprocess_result(
            ["--- Page 1 ---", "Page 1 text", "[IMAGE:0]",
             "--- Page 2 ---", "Text only page",
             "--- Page 3 ---", "Page 3 text", "[IMAGE:1]"],
            images=mock_images,
        )

        tool = ReadFileTool()
        with patch("src.tools.file_operations.subprocess.run", return_value=mock_result):
            result = tool.execute(file_path=str(pdf_path), extract_images=True)

        assert isinstance(result.output, list)
        image_blocks = [b for b in result.output if b.get("type") == "image_url"]
        assert len(image_blocks) == 2
        assert result.metadata["image_count"] == 2


# ---------------------------------------------------------------------------
# _frame_tool_result multimodal tests
# ---------------------------------------------------------------------------


class TestFrameToolResultMultimodal:
    """Tests for _frame_tool_result handling both str and list content."""

    def test_string_content_unchanged(self):
        """Test that string content produces the same framing as before."""
        from src.core.agent import _frame_tool_result

        result = _frame_tool_result("hello world", "read_file")

        assert isinstance(result, str)
        assert "[TOOL OUTPUT from read_file" in result
        assert "hello world" in result
        assert "[END TOOL OUTPUT]" in result

    def test_list_content_frames_text_blocks(self):
        """Test that list content frames text blocks and passes images through."""
        from src.core.agent import _frame_tool_result

        content = [
            {"type": "text", "text": "Document text here"},
            {"type": "image_url", "image_url": {"url": "data:image/png;base64,abc"}},
            {"type": "text", "text": "1 embedded image(s)"},
        ]
        result = _frame_tool_result(content, "read_file")

        assert isinstance(result, list)
        assert len(result) == 3

        # Text blocks should be framed
        assert result[0]["type"] == "text"
        assert "[TOOL OUTPUT from read_file" in result[0]["text"]
        assert "Document text here" in result[0]["text"]
        assert "[END TOOL OUTPUT]" in result[0]["text"]

        # Image blocks should pass through unchanged
        assert result[1]["type"] == "image_url"
        assert result[1]["image_url"]["url"] == "data:image/png;base64,abc"

        # Second text block also framed
        assert "[TOOL OUTPUT from read_file" in result[2]["text"]

    def test_empty_list_returns_empty_list(self):
        """Test that empty list content returns empty list."""
        from src.core.agent import _frame_tool_result

        result = _frame_tool_result([], "test_tool")
        assert result == []


# ---------------------------------------------------------------------------
# Anthropic backend multimodal tool result tests
# ---------------------------------------------------------------------------


class TestAnthropicMultimodalToolResult:
    """Tests for Anthropic backend handling multimodal tool results."""

    @pytest.fixture
    def backend(self):
        """Create an AnthropicBackend for translation testing."""
        with patch("src.llm.anthropic_backend.Anthropic"), \
             patch("src.llm.anthropic_backend.AsyncAnthropic"):
            from src.llm.anthropic_backend import AnthropicBackend
            config = LLMConfig(
                backend_type=LLMBackendType.ANTHROPIC,
                model_name="claude-sonnet-4-5-20250929",
                base_url="https://api.anthropic.com",
                context_window=200000,
                temperature=0.2,
                max_tokens=16384,
                top_p=0.95,
            )
            return AnthropicBackend(config, api_key="test-key")

    def test_string_tool_result_unchanged(self, backend):
        """Test that string tool results translate as before."""
        messages = [
            {"role": "user", "content": "Read the file"},
            {"role": "assistant", "content": "I'll read it.", "tool_calls": [
                {"id": "tc_123", "type": "function", "function": {"name": "read_file", "arguments": "{}"}}
            ]},
            {"role": "tool", "tool_call_id": "tc_123", "content": "file contents here"},
        ]

        _system, translated = backend._translate_messages(messages)

        # Find the tool_result block
        tool_msg = [m for m in translated if m["role"] == "user" and isinstance(m.get("content"), list)
                     and any(b.get("type") == "tool_result" for b in m["content"])]
        assert len(tool_msg) >= 1
        tool_result = [b for b in tool_msg[-1]["content"] if b.get("type") == "tool_result"][0]
        assert tool_result["content"] == "file contents here"
        assert tool_result["tool_use_id"] == "tc_123"

    def test_multimodal_tool_result_converted(self, backend):
        """Test that list tool results convert image_url to Anthropic image blocks."""
        multimodal_content = [
            {"type": "text", "text": "Document text"},
            {"type": "image_url", "image_url": {"url": "data:image/png;base64,dGVzdA=="}},
            {"type": "text", "text": "1 embedded image(s)"},
        ]

        messages = [
            {"role": "user", "content": "Read the PDF"},
            {"role": "assistant", "content": "Reading.", "tool_calls": [
                {"id": "tc_456", "type": "function", "function": {"name": "read_file", "arguments": "{}"}}
            ]},
            {"role": "tool", "tool_call_id": "tc_456", "content": multimodal_content},
        ]

        _system, translated = backend._translate_messages(messages)

        # Find the tool_result
        tool_msg = [m for m in translated if m["role"] == "user" and isinstance(m.get("content"), list)
                     and any(b.get("type") == "tool_result" for b in m["content"])]
        assert len(tool_msg) >= 1
        tool_result = [b for b in tool_msg[-1]["content"] if b.get("type") == "tool_result"][0]

        # Content should be a list of Anthropic-format blocks
        assert isinstance(tool_result["content"], list)
        assert len(tool_result["content"]) == 3

        # First: text block
        assert tool_result["content"][0]["type"] == "text"
        assert tool_result["content"][0]["text"] == "Document text"

        # Second: image block (converted from image_url to Anthropic image)
        img_block = tool_result["content"][1]
        assert img_block["type"] == "image"
        assert img_block["source"]["type"] == "base64"
        assert img_block["source"]["media_type"] == "image/png"
        assert img_block["source"]["data"] == "dGVzdA=="

        # Third: text block
        assert tool_result["content"][2]["type"] == "text"


# ---------------------------------------------------------------------------
# Display path multimodal safety tests
# ---------------------------------------------------------------------------


class TestGetTextContentMultimodal:
    """Tests for Message.get_text_content() with multimodal content."""

    def test_string_content_unchanged(self):
        """get_text_content() returns string content as-is."""
        from src.session.models.message import Message

        msg = Message(role="tool", content="plain text", tool_call_id="tc_1")
        assert msg.get_text_content() == "plain text"

    def test_none_content_returns_empty(self):
        """get_text_content() returns empty string for None."""
        from src.session.models.message import Message

        msg = Message(role="tool", content=None, tool_call_id="tc_1")
        assert msg.get_text_content() == ""

    def test_list_content_extracts_text_blocks(self):
        """get_text_content() extracts text from multimodal list content."""
        from src.session.models.message import Message

        msg = Message(
            role="tool",
            content=[
                {"type": "text", "text": "Document text here"},
                {"type": "image_url", "image_url": {"url": "data:image/png;base64,abc"}},
                {"type": "text", "text": "[1 embedded image(s) extracted from document]"},
            ],
            tool_call_id="tc_1",
        )
        text = msg.get_text_content()

        assert isinstance(text, str)
        assert "Document text here" in text
        assert "1 embedded image(s) extracted from document" in text
        assert "1 image(s) embedded in document" in text  # count note from helper
        assert "data:image/png" not in text  # base64 data NOT in display text

    def test_list_content_no_images(self):
        """get_text_content() handles list content with only text blocks."""
        from src.session.models.message import Message

        msg = Message(
            role="tool",
            content=[{"type": "text", "text": "Only text"}],
            tool_call_id="tc_1",
        )
        text = msg.get_text_content()

        assert text == "Only text"
        assert "image" not in text

    def test_empty_list_returns_empty(self):
        """get_text_content() handles empty list."""
        from src.session.models.message import Message

        msg = Message(role="tool", content=[], tool_call_id="tc_1")
        assert msg.get_text_content() == ""


class TestSerializerMultimodal:
    """Tests for serializer handling of multimodal tool results."""

    def test_string_result_serialized_normally(self):
        """String results are serialized as before."""
        from src.server.serializers import serialize_store_notification
        from src.core.events import ToolStatus as CoreToolStatus
        from src.session.store.memory_store import StoreNotification, StoreEvent, ToolExecutionState

        notification = StoreNotification(
            event=StoreEvent.TOOL_STATE_UPDATED,
            tool_call_id="tc_1",
            tool_state=ToolExecutionState(
                status=CoreToolStatus.SUCCESS,
                result="plain text result",
            ),
        )

        serialized = serialize_store_notification(notification)
        assert serialized["data"]["result"] == "plain text result"

    def test_list_result_extracted_to_text(self):
        """Multimodal list results are extracted to display text."""
        from src.server.serializers import serialize_store_notification
        from src.core.events import ToolStatus as CoreToolStatus
        from src.session.store.memory_store import StoreNotification, StoreEvent, ToolExecutionState

        notification = StoreNotification(
            event=StoreEvent.TOOL_STATE_UPDATED,
            tool_call_id="tc_2",
            tool_state=ToolExecutionState(
                status=CoreToolStatus.SUCCESS,
                result=[
                    {"type": "text", "text": "Document content"},
                    {"type": "image_url", "image_url": {"url": "data:image/png;base64,abc"}},
                    {"type": "text", "text": "[1 embedded image(s)]"},
                ],
            ),
        )

        serialized = serialize_store_notification(notification)
        result_text = serialized["data"]["result"]

        assert isinstance(result_text, str)
        assert "Document content" in result_text
        assert "1 image(s) embedded in document" in result_text
        assert "data:image/png" not in result_text
