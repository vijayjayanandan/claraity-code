"""
Subprocess-safe document text and image extractor.

This module is designed to run as a child process via `python -m src.tools.document_extractor`.
It extracts text (and optionally images) from PDF and DOCX files and writes JSON to stdout.

Running in a subprocess isolates the main agent from C-level crashes in
PyMuPDF's MuPDF library. If the parser segfaults, only this child process
dies -- the agent receives a non-zero exit code and reports an error.

Scout-then-render design:
- Default (no --extract-images): text extraction + image hint lines per page/paragraph
- With --extract-images: text + page screenshots (PDF) or inline images (DOCX)
- With --pages: filter to specific pages (PDF only)

Usage:
    python -m src.tools.document_extractor <file_path> [--format pdf|docx]
        [--max-lines 10000] [--max-zip-size 209715200] [--max-zip-ratio 100]
        [--extract-images] [--max-images 20] [--max-image-bytes 1048576]
        [--max-total-image-bytes 5242880] [--pages 1-5,7]

Output (stdout): JSON object with keys:
    - "lines": list of extracted text lines (with image hints or [IMAGE:N] markers)
    - "images": list of image dicts when --extract-images (base64, media_type, source)
    - "error": error string if extraction failed, null otherwise
"""

import argparse
import base64
import json
import sys
import zipfile
from pathlib import Path


def check_zip_bomb(
    path: Path,
    max_decompressed: int = 200 * 1024 * 1024,
    max_ratio: int = 100,
) -> str | None:
    """Check DOCX ZIP for decompression bomb. Returns error string or None."""
    try:
        with zipfile.ZipFile(str(path), "r") as zf:
            total = sum(info.file_size for info in zf.infolist())
            compressed = path.stat().st_size

            if total > max_decompressed:
                return (
                    f"DOCX decompressed size too large: {total:,} bytes "
                    f"(limit: {max_decompressed:,})"
                )
            if compressed > 0 and (total / compressed) > max_ratio:
                return (
                    f"Suspicious compression ratio: {total / compressed:.0f}:1 "
                    f"(limit: {max_ratio}:1)"
                )
    except zipfile.BadZipFile:
        pass  # Let the document library handle it
    return None


def _parse_pages(pages_str: str, total_pages: int) -> set[int]:
    """Parse pages specification into a set of 0-indexed page numbers.

    Accepts: "3", "1-5", "3,7,9", "1-3,7,10-12"
    Input is 1-indexed (user-facing), output is 0-indexed (internal).
    Malformed parts are silently skipped.
    """
    result: set[int] = set()
    for part in pages_str.split(","):
        part = part.strip()
        if not part:
            continue
        try:
            if "-" in part:
                lo, hi = part.split("-", 1)
                lo_idx = max(0, int(lo.strip()) - 1)
                hi_idx = min(total_pages, int(hi.strip()))
                result.update(range(lo_idx, hi_idx))
            else:
                idx = int(part.strip()) - 1
                if 0 <= idx < total_pages:
                    result.add(idx)
        except ValueError:
            continue  # skip malformed parts gracefully
    return result


def _count_page_images(page) -> int:
    """Count embedded images on a PyMuPDF page (cheap, no extraction)."""
    try:
        return len(page.get_images(full=True))
    except Exception:
        return 0


def _extract_page_images_inline(
    doc,
    page,
    page_num: int,
    lines: list[str],
    images: list[dict],
    seen_xrefs: set[int],
    total_bytes: list[int],
    max_images: int,
    max_image_bytes: int,
    max_total_bytes: int,
) -> None:
    """Extract individual embedded images from a PDF page at original resolution.

    Inserts [IMAGE:N] markers into lines and appends image dicts to images.
    Uses xref-based extraction (not page rendering) for sharper, smaller images.
    Modifies lines, images, seen_xrefs, total_bytes in-place.
    """
    import fitz  # already imported by caller

    if len(images) >= max_images:
        return

    try:
        img_list = page.get_images(full=True)
    except Exception:
        return

    for img_info in img_list:
        if len(images) >= max_images:
            break

        xref = img_info[0]
        if xref in seen_xrefs:
            continue
        seen_xrefs.add(xref)

        try:
            pix = fitz.Pixmap(doc, xref)

            # Convert CMYK/other colorspaces to RGB
            if pix.n > 4:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            elif pix.alpha:
                pix = fitz.Pixmap(pix, 0)  # drop alpha

            png_bytes = pix.tobytes("png")

            if len(png_bytes) > max_image_bytes:
                continue  # skip oversized images
            if total_bytes[0] + len(png_bytes) > max_total_bytes:
                break  # hit total budget

            total_bytes[0] += len(png_bytes)
            lines.append(f"[IMAGE:{len(images)}]")
            images.append({
                "base64": base64.b64encode(png_bytes).decode("ascii"),
                "media_type": "image/png",
                "source": f"page_{page_num + 1}",
            })
        except Exception:
            continue


def extract_pdf(
    path: Path,
    max_lines: int,
    extract_images: bool = False,
    max_images: int = 20,
    max_image_bytes: int = 1_048_576,
    max_total_image_bytes: int = 5_242_880,
    pages_filter: set[int] | None = None,
) -> tuple[list[str], list[dict], str | None]:
    """Extract text, tables, and optionally individual images from a PDF file.

    Scout mode (extract_images=False):
        Extracts text from all pages. Inserts hint lines on pages that contain
        images, e.g. "[1 image on this page -- use extract_images=true to see]".

    Extract mode (extract_images=True):
        Extracts text from all pages. For pages with images, extracts each
        embedded image individually at original resolution (via xref) and
        inserts [IMAGE:N] markers after the page's text.
        Text-only pages are unaffected (no token overhead).

    Args:
        pages_filter: If set, only process these 0-indexed page numbers.

    Returns (lines, images, error). images is empty in scout mode.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return [], [], "PyMuPDF is required to read PDF files. Install it: pip install PyMuPDF"

    lines: list[str] = []
    images: list[dict] = []
    seen_xrefs: set[int] = set()
    total_bytes: list[int] = [0]  # mutable counter

    try:
        doc = fitz.open(str(path))
    except Exception as e:
        return [], [], f"Failed to open PDF: {e}"

    try:
        total_pages = len(doc)
        active_pages = pages_filter if pages_filter is not None else set(range(total_pages))

        for page_num in range(total_pages):
            if page_num not in active_pages:
                continue

            page = doc[page_num]

            lines.append(f"--- Page {page_num + 1} of {total_pages} ---")
            lines.append("")

            # Extract text
            text = page.get_text()
            if text.strip():
                for line in text.splitlines():
                    lines.append(line)
                    if len(lines) >= max_lines:
                        lines.append(
                            f"[... extraction stopped: exceeded {max_lines:,} lines.]"
                        )
                        return lines, images, None

            # Extract tables
            if hasattr(page, "find_tables"):
                tables = page.find_tables()
                for table_idx, table in enumerate(tables):
                    data = table.extract()
                    if data:
                        lines.append("")
                        lines.append(f"[Table {table_idx + 1}]")
                        for row in data:
                            cells = [str(c) if c is not None else "" for c in row]
                            lines.append("| " + " | ".join(cells) + " |")
                            if len(lines) >= max_lines:
                                lines.append(
                                    f"[... extraction stopped: exceeded {max_lines:,} lines.]"
                                )
                                return lines, images, None

            # Image handling: scout hints or individual extraction
            img_count = _count_page_images(page)
            if img_count > 0:
                if extract_images and len(images) < max_images:
                    # Extract mode: individual images at original resolution
                    _extract_page_images_inline(
                        doc, page, page_num, lines, images, seen_xrefs,
                        total_bytes, max_images, max_image_bytes, max_total_image_bytes,
                    )
                else:
                    # Scout mode: hint line only
                    lines.append(
                        f"[{img_count} image(s) on this page"
                        f" -- use extract_images=true to see]"
                    )

            lines.append("")
    except Exception as e:
        return [], [], f"Failed to extract PDF text: {e}"
    finally:
        doc.close()

    return lines, images, None


def _count_paragraph_images(para) -> int:
    """Count inline images in a python-docx paragraph (cheap, XML scan)."""
    try:
        blips = para._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        )
        return len(blips)
    except Exception:
        return 0


def _extract_paragraph_images_inline(
    doc,
    para,
    lines: list[str],
    images: list[dict],
    total_bytes: list[int],
    max_images: int,
    max_image_bytes: int,
    max_total_bytes: int,
) -> None:
    """Extract inline images from a DOCX paragraph, inserting [IMAGE:N] markers.

    Finds <a:blip> elements in the paragraph XML, resolves relationship IDs
    to image parts, and emits base64-encoded image data.
    """
    if len(images) >= max_images:
        return

    mime_map = {
        "image/png": "image/png",
        "image/jpeg": "image/jpeg",
        "image/gif": "image/gif",
        "image/bmp": "image/bmp",
        "image/tiff": "image/tiff",
    }

    try:
        blips = para._element.findall(
            ".//{http://schemas.openxmlformats.org/drawingml/2006/main}blip"
        )
        for blip in blips:
            if len(images) >= max_images:
                break

            r_embed = blip.get(
                "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
            )
            if not r_embed or r_embed not in doc.part.rels:
                continue

            try:
                rel = doc.part.rels[r_embed]
                if "image" not in rel.reltype:
                    continue

                blob = rel.target_part.blob
                content_type = rel.target_part.content_type
                media_type = mime_map.get(content_type)
                if not media_type:
                    continue

                if len(blob) > max_image_bytes:
                    continue
                if total_bytes[0] + len(blob) > max_total_bytes:
                    break

                total_bytes[0] += len(blob)
                lines.append(f"[IMAGE:{len(images)}]")
                images.append({
                    "base64": base64.b64encode(blob).decode("ascii"),
                    "media_type": media_type,
                    "source": f"rel_{r_embed}",
                })
            except Exception:
                continue
    except Exception:
        pass


def extract_docx(
    path: Path,
    max_lines: int,
    max_zip_size: int,
    max_zip_ratio: int,
    extract_images: bool = False,
    max_images: int = 20,
    max_image_bytes: int = 1_048_576,
    max_total_image_bytes: int = 5_242_880,
) -> tuple[list[str], list[dict], str | None]:
    """Extract text, tables, and optionally inline images from a DOCX file.

    Scout mode (extract_images=False):
        Extracts text from all paragraphs/tables. Inserts hint lines after
        paragraphs containing images.

    Render mode (extract_images=True):
        Extracts text and inline images. Images are extracted at their
        paragraph position with [IMAGE:N] markers for interleaving.

    Returns (lines, images, error). images is empty in scout mode.
    """
    # Zip bomb guard
    bomb_error = check_zip_bomb(path, max_zip_size, max_zip_ratio)
    if bomb_error:
        return [], [], bomb_error

    try:
        from docx import Document
        from docx.table import Table
        from docx.text.paragraph import Paragraph
    except ImportError:
        return [], [], "python-docx is required to read Word files. Install it: pip install python-docx"

    try:
        doc = Document(str(path))
    except Exception as e:
        return [], [], f"Failed to open DOCX: {e}"

    lines: list[str] = []
    images: list[dict] = []
    total_bytes: list[int] = [0]
    table_idx = 0

    para_map: dict[int, Paragraph] = {id(p._element): p for p in doc.paragraphs}
    table_map: dict[int, Table] = {id(t._element): t for t in doc.tables}

    for element in doc.element.body:
        elem_id = id(element)

        if elem_id in para_map:
            para = para_map[elem_id]
            lines.append(para.text)
            if len(lines) >= max_lines:
                lines.append(
                    f"[... extraction stopped: exceeded {max_lines:,} lines.]"
                )
                return lines, images, None

            # Image handling: scout hints or inline extraction
            img_count = _count_paragraph_images(para)
            if img_count > 0:
                if extract_images and len(images) < max_images:
                    _extract_paragraph_images_inline(
                        doc, para, lines, images, total_bytes,
                        max_images, max_image_bytes, max_total_image_bytes,
                    )
                else:
                    lines.append(
                        f"[{img_count} image(s) in this paragraph"
                        f" -- use extract_images=true to see]"
                    )

        elif elem_id in table_map:
            table_idx += 1
            table = table_map[elem_id]
            lines.append("")
            lines.append(f"[Table {table_idx}]")
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if len(lines) >= max_lines:
                    lines.append(
                        f"[... extraction stopped: exceeded {max_lines:,} lines.]"
                    )
                    return lines, images, None
            lines.append("")

    return lines, images, None


def main() -> None:
    import os

    parser = argparse.ArgumentParser(description="Extract text and images from PDF/DOCX files")
    parser.add_argument("file_path", help="Path to the document")
    parser.add_argument("--format", choices=["pdf", "docx"], help="Document format (auto-detected from extension if omitted)")
    parser.add_argument("--max-lines", type=int, default=10_000, help="Max lines to extract")
    parser.add_argument("--max-zip-size", type=int, default=200 * 1024 * 1024, help="Max DOCX decompressed size")
    parser.add_argument("--max-zip-ratio", type=int, default=100, help="Max DOCX compression ratio")
    parser.add_argument("--extract-images", action="store_true", help="Render pages with images (PDF) or extract inline images (DOCX)")
    parser.add_argument("--max-images", type=int, default=20, help="Max images to extract")
    parser.add_argument("--max-image-bytes", type=int, default=1_048_576, help="Max bytes per image (1MB default)")
    parser.add_argument("--max-total-image-bytes", type=int, default=5_242_880, help="Max total image bytes (5MB default)")
    parser.add_argument("--pages", type=str, default=None, help="PDF pages to process: '3', '1-5', '3,7,9'")
    args = parser.parse_args()

    # Suppress library messages (e.g. PyMuPDF's "Consider using pymupdf_layout")
    # that print to stdout and pollute the JSON output.
    real_stdout = sys.stdout
    devnull_out = open(os.devnull, "w")
    devnull_err = open(os.devnull, "w")
    sys.stdout = devnull_out
    sys.stderr = devnull_err

    try:
        path = Path(args.file_path)
        if not path.is_file():
            result = {"lines": [], "images": [], "error": f"File not found: {args.file_path}"}
            sys.stdout = real_stdout
            json.dump(result, real_stdout)
            return

        fmt = args.format or path.suffix.lower().lstrip(".")

        if fmt == "pdf":
            # Parse pages filter with a generous upper bound; extract_pdf clips
            # to actual page count via its active_pages intersection.
            pages_filter = None
            if args.pages:
                pages_filter = _parse_pages(args.pages, 10_000)

            lines, images, error = extract_pdf(
                path, args.max_lines,
                extract_images=args.extract_images,
                max_images=args.max_images,
                max_image_bytes=args.max_image_bytes,
                max_total_image_bytes=args.max_total_image_bytes,
                pages_filter=pages_filter,
            )
        elif fmt == "docx":
            lines, images, error = extract_docx(
                path, args.max_lines, args.max_zip_size, args.max_zip_ratio,
                extract_images=args.extract_images,
                max_images=args.max_images,
                max_image_bytes=args.max_image_bytes,
                max_total_image_bytes=args.max_total_image_bytes,
            )
        else:
            lines, images, error = [], [], f"Unsupported format: {fmt}"

        result = {"lines": lines, "images": images, "error": error}
    finally:
        sys.stdout = real_stdout
        devnull_out.close()
        devnull_err.close()

    json.dump(result, real_stdout)


if __name__ == "__main__":
    main()
